import os
import zipfile
import re
import fitz  # PyMuPDF
from docx import Document
import time
import unicodedata

def clean_name(filename, texto_remover=""):
    """Clean filename by removing accents, spaces, and weird characters. 
    Also removes user-specified prefixes."""
    new_name = filename
    if texto_remover:
        new_name = new_name.replace(texto_remover, '')
    
    new_name = new_name.replace(' ', '_')
    new_name = unicodedata.normalize('NFKD', new_name).encode('ASCII', 'ignore').decode('ASCII')
    new_name = re.sub(r'[^a-zA-Z0-9._-]', '', new_name)
    return new_name

def extract_pdf_last_pages(filepath):
    """Reads PDF and gets text from last 4 pages."""
    text = ""
    try:
        doc = fitz.open(filepath)
        total = len(doc)
        start_page = max(0, total - 4)
        for i in range(start_page, total):
            page = doc[i]
            text += page.get_text() + " "
    except Exception as e:
        print(f"Error PDF {filepath}: {e}")
    return text

def extract_docx(filepath):
    """Reads DOCX and extracts all paragraphs text."""
    text = ""
    try:
        doc = Document(filepath)
        for p in doc.paragraphs:
            text += p.text + " "
    except Exception as e:
        pass
    return text

def parse_pages(text):
    """Finds 'p. 10-15' or 'pp. 10-15' on the last 2000 chars of text."""
    # Normalize spaces
    text = re.sub(r'\s+', ' ', text.replace('\n', ' '))
    text = text[-2000:]
    
    # Matches like pp. 50 - 55 ou p. 50-55
    matches = re.findall(r'(?:p\.|pp\.)\s*[0-9]{1,5}\s*[-–—]\s*[0-9]{1,5}', text)
    if matches:
        last_match = matches[-1]
        # Keep only numbers and dashes
        pag_normalizada = re.sub(r'[^0-9]+', '-', last_match).strip('-')
        return pag_normalizada
    return None

def estrutural_validar_docx(caminho_arquivo):
    """Valida a estrutura do docx (questões e gabaritos) igual ao validar_questoes.py original"""
    try:
        doc = Document(caminho_arquivo)
    except Exception as e:
        return False, "Erro crítico ao abrir", [str(e)]

    erros = []
    
    ultima_questao = 0
    questoes_encontradas = 0
    esperando_alternativa = 'a'
    ultimo_gabarito = 0
    gabaritos_encontrados = 0
    
    regex_numero = re.compile(r"^\s*(\d+)")
    regex_letra = re.compile(r"^\s*([a-z])")
    regex_gabarito_num = re.compile(r"Atividade\s+(\d+)", re.IGNORECASE)

    for i, p in enumerate(doc.paragraphs):
        texto = p.text.strip()
        estilo = p.style.name

        if estilo == "3c Atividade Enunciado":
            match = regex_numero.match(texto)
            if match:
                num_atual = int(match.group(1))
                questoes_encontradas += 1
                
                if num_atual != ultima_questao + 1:
                    erros.append(f"Q. fora de ordem. Esp: {ultima_questao + 1}, Enc: {num_atual}")
                ultima_questao = num_atual
                esperando_alternativa = 'a'
            else:
                erros.append(f"Enunciado sem número: '{texto[:15]}...'")

        elif estilo == "3b Atividade alternativa":
            match = regex_letra.match(texto)
            if match:
                letra_atual = match.group(1)
                if letra_atual != esperando_alternativa:
                    erros.append(f"Alt errada Q.{ultima_questao}. Esp: '{esperando_alternativa}', Enc: '{letra_atual}'")
                esperando_alternativa = chr(ord(letra_atual) + 1)
            else:
                erros.append(f"Alternativa sem letra: '{texto[:15]}...'")

        elif estilo == "3c Resposta gabarito":
            match = regex_gabarito_num.search(texto)
            if match:
                num_gabarito = int(match.group(1))
                gabaritos_encontrados += 1
                
                if num_gabarito != ultimo_gabarito + 1:
                    erros.append(f"Gab fora de ordem. Esp: {ultimo_gabarito + 1}, Enc: {num_gabarito}")
                ultimo_gabarito = num_gabarito

    if questoes_encontradas != gabaritos_encontrados:
        erros.append(f"Qtd: {questoes_encontradas} quest/ {gabaritos_encontrados} gab")

    if not erros:
        return True, f"Questões: {questoes_encontradas}", []
    else:
        return False, f"{len(erros)} erros", erros


def process_files(job_dir, file_paths, texto_remover, q):
    """Background task that processes the files and puts logs in Queue."""
    def emit(msg_type, content):
        q.put({"type": msg_type, "msg": content})
        time.sleep(0.12) # Delay for nice boot effect

    emit("log", "[ INIT ] Inicializando sistema de validação Prepub")
    time.sleep(0.5)

    # 1. Unzip
    zips = [f for f in file_paths if f.lower().endswith('.zip')]
    if zips:
        emit("log", f"[ OK ] {len(zips)} pacotes ZIP localizados. Descompactando payload...")
        for z in zips:
            try:
                with zipfile.ZipFile(z, 'r') as zip_ref:
                    # prevent nested dirs? just extract all using namelist and write to job_dir if we wanted flat
                    # simple extract:
                    zip_ref.extractall(job_dir)
            except Exception as e:
                emit("log", f"[ ERROR ] Falha de descompactação no cluster: {os.path.basename(z)}")
    
    # Find all PDFs and DOCXs
    all_files = []
    for root, dirs, files in os.walk(job_dir):
        for f in files:
            if f.lower().endswith(('.pdf', '.docx')):
                all_files.append(os.path.join(root, f))
    
    if not all_files:
        emit("log", "[ WAIT ] Sinal vazio. Nenhum arquivo .pdf ou .docx encontrado na matriz.")
        emit("done", {"results": [], "sem_padrao": []})
        return

    emit("log", f"[ OK ] {len(all_files)} arquivos injetados no sistema. Acionando varredores quanticos...")
    time.sleep(0.5)

    mapa_paginas = {}
    sem_padrao = []
    
    total = len(all_files)
    for i, file in enumerate(sorted(all_files)):
        pct = int((i + 1) / total * 100)
        basename = os.path.basename(file)
        
        # Limpar Nome
        clean_basename = clean_name(basename, texto_remover=texto_remover)
        if clean_basename != basename:
            new_path = os.path.join(os.path.dirname(file), clean_basename)
            os.rename(file, new_path)
            file = new_path
            basename = clean_basename

        short_name = basename if len(basename) <= 30 else basename[:27] + "..."
        # Ex: [ LOAD ] 45% -> Arquivo_exemplo...
        emit("log", f"[{pct:3}%] Processando/Validando: {short_name}")

        tipo = ""
        texto = ""
        estrutura_ok = True
        estrutura_msg = ""
        estrutura_erros = []
        
        if basename.lower().endswith('.pdf'):
            tipo = "PDF"
            texto = extract_pdf_last_pages(file)
        elif basename.lower().endswith('.docx'):
            tipo = "DOCX"
            texto = extract_docx(file)
            
            # Executa a nova validacao estrutural no DOCX importada do validar_questoes
            ok, msg, errs = estrutural_validar_docx(file)
            estrutura_ok = ok
            estrutura_msg = msg
            estrutura_erros = errs
        
        pag = parse_pages(texto)
        if pag:
            if pag not in mapa_paginas:
                mapa_paginas[pag] = {"docs": [], "pdfs": []}
            if tipo == "PDF":
                mapa_paginas[pag]["pdfs"].append(basename)
            else:
                # Mantem o log de erro atrelado ao DOCX
                mapa_paginas[pag]["docs"].append({
                    "name": basename,
                    "struct_ok": estrutura_ok,
                    "struct_msg": estrutura_msg,
                    "struct_errors": estrutura_erros
                })
        else:
            sem_padrao.append(basename)

    emit("log", "[ OK ] Varredura completa! Consolidando hashes e relatorios...")
    time.sleep(0.8)

    # Compile results
    results = []
    def sort_key(pag_str):
        # extrair apenas o primeiro numero para sortear e.g "10-15" -> 10
        nums = re.findall(r'\d+', pag_str)
        return int(nums[0]) if nums else 0

    sorted_pags = sorted(mapa_paginas.keys(), key=sort_key)
    
    for pag in sorted_pags:
        data = mapa_paginas[pag]
        c_pdf = len(data["pdfs"])
        c_doc = len(data["docs"])
        
        if c_pdf >= 1 and c_doc >= 1:
            if c_pdf > 1 or c_doc > 1:
                status = "DUPLO"
            else:
                status = "OK"
        else:
            status = "FALHOU"
        
        results.append({
            "pag": pag,
            "status": status,
            "pdfs": data["pdfs"],
            "docs": data["docs"]  # agora eh lista de objs contendo {name, struct_ok, struct_errors}
        })
    
    emit("done", {
        "results": results,
        "sem_padrao": sem_padrao
    })
